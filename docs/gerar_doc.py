"""
Gerador de documentação de gerenciamento de projetos para VestWeb
Segue o framework: Escopo, Tempo, RH, Comunicação, Risco, Aquisições, Duração
x fases: Conceito/Definição, Planejamento, Execução, Fechamento
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Margens ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# --- Estilo padrão ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='AAAAAA'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_cell_text(cell, text, bold=False, font_size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_bullet_items(cell, items, font_size=8.5, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first = True
    for item in items:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.2)
        run = p.add_run(f'• {item}')
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))

# ===================== TÍTULO PRINCIPAL =====================
title = doc.add_heading('', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Documentação de Gerenciamento de Projetos')
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('VestWeb — Plataforma de Estudos para Vestibular')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
r.font.italic = True

doc.add_paragraph()

# ===================== SEÇÃO 1: VISÃO GERAL DO PROJETO =====================
h = doc.add_heading('1. Visão Geral do Projeto', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

overview_data = [
    ('Nome do Projeto', 'VestWeb'),
    ('Tipo', 'Plataforma SaaS de Estudos para Vestibular (Web)'),
    ('Responsável', 'Luis Filipe'),
    ('Data de Início', '2024'),
    ('Status', 'Em produção (deploy ativo)'),
    ('Repositório', 'c:\\projects\\VestWeb'),
]

tbl = doc.add_table(rows=len(overview_data), cols=2)
tbl.style = 'Table Grid'
for i, (label, value) in enumerate(overview_data):
    row = tbl.rows[i]
    set_cell_bg(row.cells[0], 'E8EAF6')
    set_cell_bg(row.cells[1], 'FFFFFF')
    add_cell_text(row.cells[0], label, bold=True, font_size=9, color='1A235A')
    add_cell_text(row.cells[1], value, font_size=9)

doc.add_paragraph()

# Descrição
desc = doc.add_paragraph()
desc.add_run('Descrição: ').bold = True
desc.add_run(
    'O VestWeb é uma plataforma web SaaS voltada para estudantes que se preparam para vestibulares (ENEM e outros). '
    'Oferece banco de questões com mais de 35 modelos de dados, simulados personalizados, flashcards, correção de redação por IA, '
    'mentoria ao vivo, gamificação (pontos, badges, streak), comunidade, vídeo-aulas (Sinaflix), calendário de revisão e área do professor.'
)
desc.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ===================== SEÇÃO 2: STACK TECNOLÓGICA =====================
h2 = doc.add_heading('2. Stack Tecnológica', level=1)
h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

stack_data = [
    ('Frontend', 'React 18 + Vite + TypeScript + Tailwind CSS', 'client/'),
    ('Backend', 'Node.js + Express + Sequelize ORM', 'server/'),
    ('ENEM API', 'Next.js + Prisma (serviço separado)', 'enem-api/'),
    ('Banco de Dados', 'PostgreSQL (Supabase em produção)', '—'),
    ('Autenticação', 'JWT (tokens em localStorage)', '—'),
    ('Pagamentos', 'Stripe (webhook + checkout sessions)', '—'),
    ('Deploy Frontend', 'Vercel', 'client/'),
    ('Deploy Backend', 'Railway', 'server/'),
    ('Infra BD', 'Supabase', '—'),
]

tbl2 = doc.add_table(rows=len(stack_data)+1, cols=3)
tbl2.style = 'Table Grid'
headers = ['Camada', 'Tecnologia', 'Pasta']
header_colors = ['1A235A', '1A235A', '1A235A']
for j, h_txt in enumerate(headers):
    set_cell_bg(tbl2.rows[0].cells[j], '1A235A')
    add_cell_text(tbl2.rows[0].cells[j], h_txt, bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (camada, tech, pasta) in enumerate(stack_data):
    row = tbl2.rows[i+1]
    bg = 'F5F5FF' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    set_cell_bg(row.cells[2], bg)
    add_cell_text(row.cells[0], camada, bold=True, font_size=9, color='1A235A')
    add_cell_text(row.cells[1], tech, font_size=9)
    add_cell_text(row.cells[2], pasta, font_size=9, italic=True)

doc.add_paragraph()

# ===================== SEÇÃO 3: MÓDULOS DO SISTEMA =====================
h3 = doc.add_heading('3. Módulos do Sistema', level=1)
h3.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

modulos = [
    ('Banco de Questões', 'Questões ENEM com filtros por matéria, vestibular, dificuldade (proxy por ano). Toggle Praticar/Banco com grade de status.'),
    ('Simulados', 'Geração automática de simulados por regras (sem IA paga). Seleção aleatória com ORDER BY RANDOM(). Limite de 50 questões.'),
    ('Flashcards', 'Gerados no frontend a partir de questões. Frente = enunciado, Verso = texto da alternativa correta.'),
    ('Correção de Redação', 'Envio de redação para correção via IA. Rota POST /api/essay.'),
    ('Mentoria', 'Agendamento de sessões de mentoria ao vivo com mentores cadastrados.'),
    ('Sinaflix (Vídeos)', 'Plataforma interna de vídeo-aulas com progresso por vídeo, favoritos e área do professor.'),
    ('Gamificação', 'Sistema de pontos, badges, streak diário, ranking entre alunos.'),
    ('Comunidade', 'Posts, comentários, likes, dúvidas de alunos (StudentDoubt), relatórios.'),
    ('Calendário de Revisão', 'Calendário de eventos de estudo (StudyEvent) para organização pessoal.'),
    ('Área do Professor', 'Portal separado com login próprio. Gestão de sessões, vídeos, questões e alunos.'),
    ('Área Admin', 'Painel administrativo para gestão da plataforma.'),
    ('Pagamentos', 'Assinatura via Stripe. Webhook configurado. Modelos: Subscription, PaymentCancel, PaymentSuccess.'),
    ('Autenticação', 'Login separado para aluno e professor. JWT com 403 para acesso cruzado. Erros distintos: 404 (não encontrado), 401 (senha errada).'),
]

tbl3 = doc.add_table(rows=len(modulos)+1, cols=2)
tbl3.style = 'Table Grid'
set_cell_bg(tbl3.rows[0].cells[0], '1A235A')
set_cell_bg(tbl3.rows[0].cells[1], '1A235A')
add_cell_text(tbl3.rows[0].cells[0], 'Módulo', bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tbl3.rows[0].cells[1], 'Descrição', bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (mod, desc) in enumerate(modulos):
    row = tbl3.rows[i+1]
    bg = 'EEF2FF' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    add_cell_text(row.cells[0], mod, bold=True, font_size=9, color='1A235A')
    add_cell_text(row.cells[1], desc, font_size=9)

doc.add_paragraph()

# ===================== SEÇÃO 4: GERENCIAMENTO DE PROJETOS =====================
h4 = doc.add_heading('4. Gerenciamento de Projetos — Framework PMI', level=1)
h4.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

sub4 = doc.add_paragraph(
    'A tabela abaixo apresenta as atividades do projeto VestWeb mapeadas nas fases e áreas do framework básico de gerenciamento de projetos.'
)
sub4.runs[0].font.size = Pt(10)
doc.add_paragraph()

# Cores das colunas (fases)
FASE_COLORS = {
    'Fase / Área':       ('F5F5F5', '333333'),
    'Conceito/Definição': ('FFE0B2', '5D4037'),
    'Planejamento':       ('E3F2FD', '0D47A1'),
    'Execução':           ('E8F5E9', '1B5E20'),
    'Fechamento':         ('FCE4EC', '880E4F'),
}

# Dados da tabela de gerenciamento
pm_data = [
    # (Área, cor_area_hex, [conceito], [planejamento], [execução], [fechamento])
    (
        'Escopo',
        'FFF3E0',
        [
            'Definir MVP: banco de questões + simulados + autenticação JWT',
            'Conduzir entrevistas com estudantes e professores',
            'Criar declaração de alto nível: plataforma SaaS de estudos vestibulares',
            'Levantamento de requisitos: questões ENEM, simulados, flashcards, gamificação, mentoria',
        ],
        [
            'Solidificar declaração de escopo (plataforma multi-módulo)',
            'Documentar módulos: Questões, Simulados, Flashcards, Redação, Mentoria, Vídeos, Gamificação, Comunidade',
            'Estabelecer plano de controle de mudanças via issues no GitHub',
            'Definir backlog priorizado por valor para o estudante',
        ],
        [
            'Verificar plano: todos os módulos implementados conforme escopo',
            'Controlar escopo: evitar feature creep (novas features apenas via backlog aprovado)',
            'Validar integração entre módulos (auth → questões → simulados → gamificação)',
            'Testes de aceitação por tipo de usuário (aluno, professor, admin)',
        ],
        [
            'Garantir que escopo atende às necessidades de alunos e professores',
            'Aceitação do cliente: plataforma funcional em produção (Vercel + Railway + Supabase)',
            'Documentar versão entregue e módulos ativos',
            'Validar que todos os fluxos críticos funcionam em produção',
        ],
    ),
    (
        'Tempo',
        'FFF8E1',
        [
            'Estimar orçamento de tempo por módulo (backend + frontend)',
            'Identificar custos fixos: Supabase, Railway, Vercel (planos gratuitos/pagos)',
            'Priorizar entregas: autenticação → questões → simulados → features extras',
        ],
        [
            'Identificar especificações dos produtos: React 18, Node.js, PostgreSQL, Stripe',
            'Definir critérios de medição: uptime > 99%, resposta API < 500ms',
            'Planejar sprints por módulo com estimativas de horas',
            'Estabelecer janelas de deploy (evitar horário de pico)',
        ],
        [
            'Gerenciar orçamento de infra (Railway + Supabase — monitorar uso)',
            'Aplicar ações corretivas em bugs críticos (ex: json_agg retornando [null])',
            'Monitorar tempos de resposta das queries PostgreSQL',
            'Ajustar pool de conexões Sequelize (max: 5) conforme carga',
        ],
        [
            'Realocar recursos de infra conforme crescimento de usuários',
            'Revisar desempenho: queries lentas, endpoints com alta latência',
            'Documentar métricas de desempenho da versão entregue',
        ],
    ),
    (
        'RH',
        'F3E5F5',
        [
            'Definir expectativas: Dev Full-Stack responsável por frontend, backend e infra',
            'Mapear habilidades necessárias: React/TS, Node.js/Express, PostgreSQL, Stripe, Docker',
            'Identificar gaps: design UI/UX, marketing, suporte ao aluno',
        ],
        [
            'Aprovar e definir papéis: Desenvolvedor (Luis Filipe), Professor (usuário-tipo), Admin',
            'Contratar equipe de suporte/tutores conforme crescimento',
            'Criar plano de comunicação interno (GitHub Issues + commits descritivos)',
            'Documentar onboarding para novos colaboradores',
        ],
        [
            'Executar tarefas de desenvolvimento seguindo backlog priorizado',
            'Gerenciar riscos técnicos: dual ORM (Sequelize + Prisma), migração de dados',
            'Capacitar professores no portal específico (TeacherArea)',
            'Revisar PRs e manter qualidade de código',
        ],
        [
            'Realocar recursos: definir papéis de manutenção pós-lançamento',
            'Revisar desempenho da equipe e processos de desenvolvimento',
            'Documentar lições aprendidas para próximas versões',
        ],
    ),
    (
        'Comunicação',
        'E8F5E9',
        [
            'Definir canais de comunicação com stakeholders (alunos, professores)',
            'Estabelecer formato de feedback de usuários (email: contato@bratloware.com)',
            'Definir política de comunicação de erros (mensagens de erro distintas por tipo)',
        ],
        [
            'Aprovar plano de comunicação: GitHub para dev, email para suporte',
            'Criar plano de aquisições de ferramentas (Supabase, Railway, Vercel, Stripe)',
            'Documentar APIs públicas e internas (rotas, payloads, autenticação)',
            'Planejar comunicação de manutenção e downtime para usuários',
        ],
        [
            'Executar comunicação: notificações in-app, anúncios no dashboard (Announcements)',
            'Gerenciar com análise de risco: monitorar webhooks Stripe, falhas de autenticação',
            'Manter logs de requisições (requestLoggerMiddleware) para auditoria',
            'Responder dúvidas de alunos via módulo StudentDoubt',
        ],
        [
            'Realocar recursos de comunicação (suporte pós-lançamento)',
            'Revisar desempenho: taxa de resolução de dúvidas, NPS de alunos',
            'Documentar processos de comunicação para a próxima versão',
        ],
    ),
    (
        'Risco',
        'E1F5FE',
        [
            'Identificar habilidades requeridas: React, Node.js, PostgreSQL, Stripe, JWT, Docker',
            'Mapear riscos técnicos: dual ORM (Sequelize + Prisma), pool de conexões, CORS',
            'Riscos de negócio: dependência de APIs externas (ENEM API separada)',
            'Risco de dados: questões importadas via Prisma, app usa Sequelize',
        ],
        [
            'Monitorar riscos de infraestrutura (limites de plano Railway/Supabase)',
            'Plano de mitigação: DATABASE_URL com fallback para variáveis individuais',
            'Plano de backup: Supabase com backups automáticos',
            'Documentar riscos de segurança: rate limiting, validação de schemas (Joi/Zod)',
        ],
        [
            'Validar orçamento de infra mensal (Railway + Supabase)',
            'Monitorar falhas de webhook Stripe (express.raw() antes de express.json())',
            'Aplicar rate limiting (rateLimitMiddleware) contra abusos',
            'Validar SSL automático em produção (Supabase connection)',
            'Monitorar erros de autenticação: 401, 403, 404 distintos',
        ],
        [
            'Obter aprovação final: deploy validado em produção',
            'Documentar incidentes e soluções aplicadas',
            'Revisar postura de segurança: JWT, bcrypt, CORS, rate limit',
        ],
    ),
    (
        'Aquisições',
        'F9FBE7',
        [
            'Identificar fornecedores: Vercel (frontend), Railway (backend), Supabase (BD), Stripe (pagamentos)',
            'Avaliar planos gratuitos x pagos de cada serviço',
            'Definir relacionamentos: contrato com Stripe para processamento de pagamentos',
        ],
        [
            'Criar plano de aquisições de serviços cloud (Vercel Hobby → Pro, Railway Starter, Supabase Free)',
            'Negociar limites de uso: pool Sequelize max 5, uploads via memoryStorage (sem disco)',
            'Documentar variáveis de ambiente necessárias (.env.example)',
            'Planejar escalabilidade: quando migrar para planos pagos',
        ],
        [
            'Validar orçamento mensal de serviços contratados',
            'Monitorar uso de banda Railway e storage Supabase',
            'Gerenciar renovação de planos e alertas de limite',
        ],
        [
            'Obter instalação: todos os serviços configurados e funcionando em produção',
            'Validar deploy completo: Vercel (client/) + Railway (server/) + Supabase',
            'Documentar configurações de produção para manutenção futura',
        ],
    ),
    (
        'Duração',
        'F5F5F5',
        [
            'Fase inicial de conceito: alguns dias (definição de stack e escopo)',
        ],
        [
            'Planejamento por módulo: algumas semanas a um mês por módulo principal',
        ],
        [
            'Desenvolvimento ativo: meses (implementação de todos os módulos)',
            'Ciclos de sprint: alguns dias a semanas por feature',
        ],
        [
            'Fechamento e deploy: algumas horas a dias por ambiente',
            'Validação em produção: 1-2 dias por release',
        ],
    ),
]

# Cores das fases
COL_HEADERS = ['Fase / Área', 'Conceito / Definição', 'Planejamento', 'Execução', 'Fechamento']
COL_BG      = ['E0E0E0',     'FFE0B2',               'BBDEFB',       'C8E6C9',   'F8BBD0']
COL_TXT     = ['212121',     '4E342E',               '0D47A1',       '1B5E20',   '880E4F']

num_rows = 1 + len(pm_data)
num_cols = 5
tbl4 = doc.add_table(rows=num_rows, cols=num_cols)
tbl4.style = 'Table Grid'

# Larguras de coluna
widths = [Cm(2.8), Cm(4.2), Cm(4.2), Cm(4.2), Cm(4.2)]

# Cabeçalho
for j, (hdr, bg, txt) in enumerate(zip(COL_HEADERS, COL_BG, COL_TXT)):
    cell = tbl4.rows[0].cells[j]
    set_cell_bg(cell, bg)
    set_cell_borders(cell, '888888')
    add_cell_text(cell, hdr, bold=True, font_size=9, color=txt, align=WD_ALIGN_PARAGRAPH.CENTER)

# Dados
for i, (area, area_bg, conceito, planejamento, execucao, fechamento) in enumerate(pm_data):
    row = tbl4.rows[i+1]
    cols_data = [conceito, planejamento, execucao, fechamento]

    # Área (col 0)
    set_cell_bg(row.cells[0], area_bg)
    set_cell_borders(row.cells[0], '888888')
    add_cell_text(row.cells[0], area, bold=True, font_size=9, color='1A235A', align=WD_ALIGN_PARAGRAPH.CENTER)

    # Fases (cols 1-4)
    phase_bgs = ['FFF8F0', 'F0F8FF', 'F0FFF4', 'FFF0F5']
    for j, (items, pbg) in enumerate(zip(cols_data, phase_bgs)):
        cell = row.cells[j+1]
        set_cell_bg(cell, pbg)
        set_cell_borders(cell, '888888')
        add_bullet_items(cell, items, font_size=8)

doc.add_paragraph()

# ===================== SEÇÃO 5: ARQUITETURA DE DEPLOY =====================
h5 = doc.add_heading('5. Arquitetura de Deploy', level=1)
h5.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

deploy_data = [
    ('Frontend (Vercel)', 'client/', 'vercel.json', 'Build multi-stage React + Nginx. SPA routing. Proxy /api/ → Railway.'),
    ('Backend (Railway)', 'server/', 'railway.toml', 'Node 20 Alpine. Start: node app.js. Healthcheck em /health.'),
    ('Banco (Supabase)', '—', '.env DATABASE_URL', 'PostgreSQL gerenciado. SSL automático. Pool Sequelize max 5.'),
    ('Docker (local dev)', '/', 'docker-compose.yml', 'PostgreSQL + backend + frontend. Comando: docker-compose up --build.'),
]

tbl5 = doc.add_table(rows=len(deploy_data)+1, cols=4)
tbl5.style = 'Table Grid'
hdrs5 = ['Serviço', 'Pasta', 'Config', 'Detalhes']
for j, h_txt in enumerate(hdrs5):
    set_cell_bg(tbl5.rows[0].cells[j], '1A235A')
    add_cell_text(tbl5.rows[0].cells[j], h_txt, bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (svc, pasta, cfg, det) in enumerate(deploy_data):
    row = tbl5.rows[i+1]
    bg = 'EEF2FF' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
    add_cell_text(row.cells[0], svc, bold=True, font_size=9, color='1A235A')
    add_cell_text(row.cells[1], pasta, font_size=9, italic=True)
    add_cell_text(row.cells[2], cfg, font_size=9, italic=True)
    add_cell_text(row.cells[3], det, font_size=9)

doc.add_paragraph()

# ===================== SEÇÃO 6: VARIÁVEIS DE AMBIENTE =====================
h6 = doc.add_heading('6. Variáveis de Ambiente', level=1)
h6.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

env_data = [
    ('DATABASE_URL', 'Backend', 'URL completa PostgreSQL (Railway/Supabase) — prioridade sobre DB_*'),
    ('DB_HOST / DB_PORT / DB_NAME / DB_USER / DB_PASSWORD', 'Backend', 'Fallback se DATABASE_URL não definido'),
    ('CLIENT_URL', 'Backend', 'Origens CORS separadas por vírgula'),
    ('JWT_SECRET', 'Backend', 'Chave para assinar tokens JWT'),
    ('STRIPE_SECRET_KEY', 'Backend', 'Chave secreta Stripe para pagamentos'),
    ('STRIPE_WEBHOOK_SECRET', 'Backend', 'Secret para validar eventos webhook Stripe'),
    ('EMAIL_USER', 'Backend', 'bratloware@gmail.com — conta de envio de emails'),
    ('EMAIL_TO', 'Backend', 'contato@bratloware.com — destino dos emails'),
    ('VITE_API_URL', 'Frontend', 'URL base da API (vazio = usa proxy Vite em dev)'),
]

tbl6 = doc.add_table(rows=len(env_data)+1, cols=3)
tbl6.style = 'Table Grid'
for j, h_txt in enumerate(['Variável', 'Escopo', 'Descrição']):
    set_cell_bg(tbl6.rows[0].cells[j], '1A235A')
    add_cell_text(tbl6.rows[0].cells[j], h_txt, bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (var, escopo, desc) in enumerate(env_data):
    row = tbl6.rows[i+1]
    bg = 'F5F5FF' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
    add_cell_text(row.cells[0], var, font_size=8.5, italic=True, color='0D47A1')
    add_cell_text(row.cells[1], escopo, font_size=9)
    add_cell_text(row.cells[2], desc, font_size=9)

doc.add_paragraph()

# ===================== SEÇÃO 7: MODELOS DE DADOS =====================
h7 = doc.add_heading('7. Modelos de Dados (Sequelize)', level=1)
h7.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x5A)

models_groups = [
    ('Usuários', ['Student', 'Teacher', 'Mentor', 'PendingStudent']),
    ('Questões', ['Question', 'Alternative', 'Answer', 'Subject', 'Topic', 'Subtopic', 'Vestibular', 'QuestionVestibular', 'QuestionSession']),
    ('Simulados', ['Simulation', 'SimulationQuestion', 'Session', 'TeacherSession']),
    ('Gamificação', ['Points', 'Badge', 'StudentBadge', 'Streak']),
    ('Vídeos', ['Video', 'InstitutionalVideo', 'VideoProgress', 'FavoriteVideo']),
    ('Comunidade', ['Post', 'Comment', 'Like', 'Report', 'StudentDoubt']),
    ('Mentoria', ['MentoringSession']),
    ('Financeiro', ['Subscription']),
    ('Conteúdo', ['Announcement', 'Banner', 'Testimonial', 'StudyEvent']),
]

tbl7 = doc.add_table(rows=len(models_groups)+1, cols=2)
tbl7.style = 'Table Grid'
for j, h_txt in enumerate(['Grupo', 'Modelos']):
    set_cell_bg(tbl7.rows[0].cells[j], '1A235A')
    add_cell_text(tbl7.rows[0].cells[j], h_txt, bold=True, font_size=9, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (grupo, models) in enumerate(models_groups):
    row = tbl7.rows[i+1]
    bg = 'EEF2FF' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
    add_cell_text(row.cells[0], grupo, bold=True, font_size=9, color='1A235A')
    add_cell_text(row.cells[1], ', '.join(models), font_size=9)

doc.add_paragraph()

# Nota sobre Prisma
p_prisma = doc.add_paragraph()
run = p_prisma.add_run('Nota ORM Dual: ')
run.bold = True
run.font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
p_prisma.add_run(
    'As tabelas "Question", "Alternative" e "Exam" são gerenciadas pelo Prisma (enem-api/). '
    'O backend principal (Sequelize) acessa essas tabelas via raw SQL (sequelize.query) no questionsController.'
).font.size = Pt(9)

doc.add_paragraph()

# ===================== RODAPÉ =====================
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = footer_p.add_run('VestWeb — Documentação de Gerenciamento de Projetos  |  Gerado em: 2026-05-10  |  Responsável: Luis Filipe')
r1.font.size = Pt(8)
r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r1.font.italic = True

# --- Salvar ---
output_path = r'c:\projects\VestWeb\docs\documentacao_gerenciamento_vestweb.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
